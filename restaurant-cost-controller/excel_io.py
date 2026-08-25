#!/usr/bin/env python3
"""Universal shared document wrappers for the restaurant cost controller.

Supported file types:
    Excel: .xlsx, .xlsm
    CSV:   .csv
    Text:  .txt
    PDF:   .pdf  (requires PyMuPDF)
    DOCX:  .docx (requires python-docx)

Existing helpers ``read_xlsx``, ``write_xlsx``, ``convert_excel_to_records``,
``convert_records_to_excel``, ``is_excel_path``, and ``preferred_extension``
are preserved for compatibility.
"""
from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any

from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter

try:
    import fitz  # PyMuPDF: text + basic table extraction.
except ImportError:
    fitz = None

try:
    import docx  # python-docx.
except ImportError:
    docx = None

SUPPORTED_EXCEL_SUFFIXES: set[str] = {".xlsx", ".xlsm"}
SUPPORTED_CSV_SUFFIXES: set[str] = {".csv"}
SUPPORTED_TXT_SUFFIXES: set[str] = {".txt"}
SUPPORTED_PDF_SUFFIXES: set[str] = {".pdf"}
SUPPORTED_DOCX_SUFFIXES: set[str] = {".docx"}

SUPPORTED_SUFFIXES: set[str] = (
    SUPPORTED_EXCEL_SUFFIXES
    | SUPPORTED_CSV_SUFFIXES
    | SUPPORTED_TXT_SUFFIXES
    | SUPPORTED_PDF_SUFFIXES
    | SUPPORTED_DOCX_SUFFIXES
)

EXTENSION_FOR_FORMAT: dict[str, str] = {
    "excel": ".xlsx",
    "xlsx": ".xlsx",
    "xlsm": ".xlsm",
    "csv": ".csv",
    "txt": ".txt",
    "pdf": ".pdf",
    "docx": ".docx",
}

FORMAT_FROM_EXTENSION: dict[str, str] = {
    ".xlsx": "excel",
    ".xlsm": "excel",
    ".csv": "csv",
    ".txt": "txt",
    ".pdf": "pdf",
    ".docx": "docx",
}

SUPPORTED_EXPORT_FORMATS: list[str] = ["excel", "csv", "txt", "pdf", "docx"]

OPTIONAL_FORMAT_FEATURES: dict[str, list[str]] = {
    "pdf": ["fitz"],
    "docx": ["python-docx"],
}


class DocumentError(Exception):
    """Raised when a document cannot be read or written."""


class ExcelError(DocumentError):
    """Raised when an Excel file cannot be read or written."""


class OptionalImportError(DocumentError):
    """Raised when an optional dependency is required but is not installed."""


def _normalize_headers(headers: list[str]) -> list[str]:
    normalized: list[str] = []
    for idx, header in enumerate(headers):
        text = str(header) if header is not None else f"column_{idx}"
        normalized.append(text.strip() or f"column_{idx}")
    return normalized


def _guess_delimiter(text: str) -> str:
    """Choose CSV-like delimiter from first few non-empty lines."""
    lines = [line for line in text.splitlines() if line.strip()]
    candidates = ["\t", ",", ";", "|"]
    best = ","
    best_count = 0
    for delimiter in candidates:
        counts = [line.count(delimiter) for line in lines[:6]]
        if not counts:
            continue
        if min(counts) <= 0:
            continue
        consistent = all(count == counts[0] for count in counts)
        if consistent and counts[0] > best_count:
            best = delimiter
            best_count = counts[0]
    return best


def read_xlsx(path: Path, *, sheet_name: str | None = None) -> list[dict[str, Any]]:
    path = Path(path)
    if not path.exists():
        raise ExcelError(f"Missing Excel file: {path}")
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:
        raise ExcelError(f"Unable to open Excel file: {exc}") from exc

    if sheet_name and sheet_name in workbook.sheetnames:
        ws = workbook[sheet_name]
    else:
        ws = workbook.active

    if ws.max_row == 0:
        return []

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return []

    headers = _normalize_headers(
        [str(cell) if cell is not None else f"column_{idx}" for idx, cell in enumerate(rows[0])]
    )
    records: list[dict[str, Any]] = []
    for row in rows[1:]:
        record: dict[str, Any] = {}
        for idx, header in enumerate(headers):
            record[header] = row[idx] if idx < len(row) else None
        records.append(record)
    try:
        workbook.close()
    except Exception:
        pass
    return records


def write_xlsx(path: Path, records: list[dict[str, Any]], *, sheet_name: str = "Sheet1") -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    ws = workbook.active
    ws.title = str(sheet_name)[:31]
    if records:
        headers = list(records[0].keys())
        ws.append(headers)
        for record in records:
            ws.append([record.get(header) for header in headers])
        for idx, header in enumerate(headers, start=1):
            max_length = len(str(header))
            for row in records:
                value = row.get(header)
                if value is not None:
                    max_length = max(max_length, len(str(value)))
            ws.column_dimensions[get_column_letter(idx)].width = min(max_length + 2, 60)
    workbook.save(str(path))
    return path


def convert_excel_to_records(path: Path) -> list[dict[str, Any]]:
    return read_xlsx(path)


def convert_records_to_excel(records: list[dict[str, Any]], path: Path, **kwargs: Any) -> Path:
    return write_xlsx(path, records, **kwargs)


def is_excel_path(path: Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXCEL_SUFFIXES


def preferred_extension() -> str:
    return ".xlsx"


# ---- CSV helpers ----


def read_csv(path: Path) -> list[dict[str, Any]]:
    path = Path(path)
    if not path.exists():
        raise DocumentError(f"Missing CSV file: {path}")
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        records: list[dict[str, Any]] = []
        for row in reader:
            record: dict[str, Any] = {}
            for key, value in row.items():
                record[key.strip() if key else key] = value
            records.append(record)
        return records


def write_csv(path: Path, records: list[dict[str, Any]], *, fieldnames: list[str] | None = None) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if not records:
        path.write_text("", encoding="utf-8")
        return path
    fieldnames = fieldnames or list(records[0].keys())
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(records)
    return path


# ---- TXT helpers ----


def _split_txt_line(line: str, delimiter: str) -> list[str]:
    if delimiter == "\t":
        return line.split("\t")
    if delimiter == "|":
        return [part.strip() for part in line.split("|")]
    if delimiter == ";":
        return [part.strip() for part in line.split(";")]
    if delimiter == ",":
        return [part.strip() for part in line.split(",")]
    return line.split()


def read_txt(path: Path) -> list[dict[str, Any]]:
    path = Path(path)
    if not path.exists():
        raise DocumentError(f"Missing text file: {path}")
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    delimiter = _guess_delimiter(text)
    lines = [line.rstrip("\n") for line in text.splitlines() if line.strip()]
    if not lines:
        return []
    headers = _normalize_headers(_split_txt_line(lines[0], delimiter))
    records: list[dict[str, Any]] = []
    for line in lines[1:]:
        parts = _split_txt_line(line, delimiter)
        record: dict[str, Any] = {}
        for idx, header in enumerate(headers):
            record[header] = parts[idx] if idx < len(parts) else None
        records.append(record)
    return records


def write_txt(path: Path, records: list[dict[str, Any]], *, delimiter: str = "\t") -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    if not records:
        path.write_text("", encoding="utf-8")
        return path
    headers = list(records[0].keys())
    lines = [delimiter.join(headers)]
    for record in records:
        lines.append(delimiter.join(str(record.get(header, "")) for header in headers))
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


# ---- PDF helpers ----


def read_pdf(path: Path) -> list[dict[str, Any]]:
    if fitz is None:
        raise OptionalImportError(
            "PDF support requires PyMuPDF. Install it with: pip install PyMuPDF"
        )
    path = Path(path)
    if not path.exists():
        raise DocumentError(f"Missing PDF file: {path}")
    document = fitz.open(path)
    try:
        tables: list[dict[str, Any]] = []
        headers = ["page", "source_text"]
        for page_index, page in enumerate(document, start=1):
            try:
                found = page.find_tables()
                raw_tables = found.extract() if found else []
            except Exception:
                raw_tables = []
            if raw_tables:
                for table in raw_tables:
                    rows = list(table)
                    if not rows:
                        continue
                    table_headers = _normalize_headers(
                        [str(cell) if cell is not None else f"column_{i}" for i, cell in enumerate(rows[0])]
                    )
                    for row in rows[1:]:
                        record: dict[str, Any] = {
                            "page": page_index,
                            "source_text": page.get_text("text").strip()[:500],
                        }
                        for idx, header in enumerate(table_headers):
                            record[header] = row[idx] if idx < len(row) else None
                        tables.append(record)
                continue
            text = page.get_text("text").strip()
            if not text:
                continue
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if len(lines) >= 2:
                header_line = lines[0]
                delimiter = _guess_delimiter(header_line)
                parts = _split_txt_line(header_line, delimiter)
                table_headers = _normalize_headers(parts)
                first = True
                for line in lines[1:]:
                    values = _split_txt_line(line, delimiter)
                    if first and len(values) < max(1, len(table_headers) - 1):
                        continue
                    first = False
                    record: dict[str, Any] = {"page": page_index, "source_text": text[:500]}
                    for idx, header in enumerate(table_headers):
                        record[header] = values[idx] if idx < len(values) else None
                    tables.append(record)
        if tables:
            return tables
        return [{"page": 1, "source_text": document[0].get_text("text").strip()[:1000] if len(document) else ""}]
    finally:
        try:
            document.close()
        except Exception:
            pass


def write_pdf(path: Path, records: list[dict[str, Any]], *, title: str = "Export") -> Path:
    if fitz is None:
        raise OptionalImportError(
            "PDF support requires PyMuPDF. Install it with: pip install PyMuPDF"
        )
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), title, fontsize=14)
    if records:
        headers = list(records[0].keys())
        y = 100
        page.insert_text((72, y), " | ".join(headers), fontsize=10)
        y += 16
        for record in records[:200]:
            page.insert_text((72, y), " | ".join(str(record.get(header, "")) for header in headers), fontsize=9)
            y += 14
            if y > 720:
                page = document.new_page()
                y = 72
    document.save(str(path))
    document.close()
    return path


# ---- DOCX helpers ----


def read_docx(path: Path) -> list[dict[str, Any]]:
    if docx is None:
        raise OptionalImportError(
            "DOCX support requires python-docx. Install it with: pip install python-docx"
        )
    path = Path(path)
    if not path.exists():
        raise DocumentError(f"Missing DOCX file: {path}")
    document = docx.Document(path)
    tables: list[dict[str, Any]] = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            continue
        headers = _normalize_headers(rows[0])
        for row in rows[1:]:
            record: dict[str, Any] = {}
            for idx, header in enumerate(headers):
                record[header] = row[idx] if idx < len(row) else None
            tables.append(record)
    if tables:
        return tables
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return [{"text": paragraph} for paragraph in paragraphs]


def write_docx(path: Path, records: list[dict[str, Any]], *, title: str = "Export") -> Path:
    if docx is None:
        raise OptionalImportError(
            "DOCX support requires python-docx. Install it with: pip install python-docx"
        )
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = docx.Document()
    document.add_heading(title, level=1)
    if records:
        headers = list(records[0].keys())
        document.add_heading("Data", level=2)
        for record in records[:500]:
            document.add_paragraph(" | ".join(str(record.get(header, "")) for header in headers))
    document.save(str(path))
    return path


# ---- Universal API ----


def document_type(path: Path) -> str:
    suffix = Path(path).suffix.lower()
    if suffix in SUPPORTED_EXCEL_SUFFIXES:
        return "excel"
    if suffix in SUPPORTED_CSV_SUFFIXES:
        return "csv"
    if suffix in SUPPORTED_TXT_SUFFIXES:
        return "txt"
    if suffix in SUPPORTED_PDF_SUFFIXES:
        return "pdf"
    if suffix in SUPPORTED_DOCX_SUFFIXES:
        return "docx"
    return "unknown"


def is_document_path(path: Path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_SUFFIXES


def resolve_document_format(identifier: str) -> tuple[str, str]:
    requested = str(identifier or "").strip().lower()
    if not requested:
        return "excel", EXTENSION_FOR_FORMAT["excel"]
    if requested in EXTENSION_FOR_FORMAT:
        return requested, EXTENSION_FOR_FORMAT[requested]
    if requested.startswith("."):
        document_format = FORMAT_FROM_EXTENSION.get(requested.lower(), "excel")
        return document_format, requested.lower()
    return "excel", EXTENSION_FOR_FORMAT["excel"]


def read_table_as(path: Path, document_format: str) -> list[dict[str, Any]]:
    resolved_format, _ = resolve_document_format(document_format)
    if resolved_format == "excel":
        return read_xlsx(path)
    if resolved_format == "csv":
        return read_csv(path)
    if resolved_format == "txt":
        return read_txt(path)
    if resolved_format == "pdf":
        return read_pdf(path)
    if resolved_format == "docx":
        return read_docx(path)
    raise DocumentError(f"Unsupported document format: {document_format}")


def write_table_as(path: Path, records: list[dict[str, Any]], document_format: str, **kwargs: Any) -> Path:
    resolved_format, _ = resolve_document_format(document_format)
    if resolved_format == "excel":
        sheet_name = str(kwargs.get("sheet_name", "Sheet1"))
        return write_xlsx(path, records, sheet_name=sheet_name)
    if resolved_format == "csv":
        fieldnames = kwargs.get("fieldnames") or (list(records[0].keys()) if records else [])
        return write_csv(path, records, fieldnames=fieldnames)
    if resolved_format == "txt":
        delimiter = str(kwargs.get("delimiter", "\t"))
        return write_txt(path, records, delimiter=delimiter)
    if resolved_format == "pdf":
        return write_pdf(path, records, title=str(kwargs.get("title", "Export")))
    if resolved_format == "docx":
        return write_docx(path, records, title=str(kwargs.get("title", "Export")))
    raise DocumentError(f"Unsupported document format: {document_format}")


def read_table(path: Path) -> list[dict[str, Any]]:
    doc_type = document_type(path)
    if doc_type == "excel":
        return read_xlsx(path)
    if doc_type == "csv":
        return read_csv(path)
    if doc_type == "txt":
        return read_txt(path)
    if doc_type == "pdf":
        return read_pdf(path)
    if doc_type == "docx":
        return read_docx(path)
    raise DocumentError(f"Unsupported document type for: {path}")


def write_table(path: Path, records: list[dict[str, Any]], **kwargs: Any) -> Path:
    doc_type = document_type(path)
    if doc_type == "excel":
        sheet_name = str(kwargs.get("sheet_name", "Sheet1"))
        return write_xlsx(path, records, sheet_name=sheet_name)
    if doc_type == "csv":
        fieldnames = kwargs.get("fieldnames") or (list(records[0].keys()) if records else [])
        return write_csv(path, records, fieldnames=fieldnames)
    if doc_type == "txt":
        delimiter = str(kwargs.get("delimiter", "\t"))
        return write_txt(path, records, delimiter=delimiter)
    if doc_type == "pdf":
        return write_pdf(path, records, title=str(kwargs.get("title", "Export")))
    if doc_type == "docx":
        return write_docx(path, records, title=str(kwargs.get("title", "Export")))
    raise DocumentError(f"Unsupported document type for write: {path}")
